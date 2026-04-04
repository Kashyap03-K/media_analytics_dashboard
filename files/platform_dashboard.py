"""
platform_dashboard.py — Unified Media Intelligence Dashboard
TV: Raw_Data3.xlsx | Social: Social_Media_ThirdParty_Data xlsx
"""
import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import sqlite3
import json
import os
import io
import base64
import tempfile
from datetime import datetime
from config import DB_PATH
from products import PRODUCTS, get_product, get_platform as _get_plat_cfg

PLATFORMS = {
    "print": {
        "label":"Print","icon": "🗞️","table":"media_print",
        "unit_opts":["Index","Reach (Total OTS)","CCMs (Coverage)","Articles"],
        "unit_keys":["index","readership","coverage","articles"],
        "col_map":{
            "Company":"company","Article_Id":"article_id","Publication":"source",
            "Edition":"edition","Publication_Date":"publication_date","AdRate":"ad_rate",
            "Headline":"headline","Article_Length":"article_length","Genre":"genre",
            "State":"state","Zone":"zone","Program_Status":"program_status",
            "Beneficial_Coverage":"beneficial_vol","Neutral_Coverage":"neutral_vol",
            "Adverse_Coverage":"adverse_vol","Total_Coverage":"total_vol",
            "Beneficial_Articles":"beneficial_art","Neutral_Articles":"neutral_art",
            "Adverse_Articles":"adverse_art","Total_Articles":"total_articles",
            "Total_OTS_in_000s":"total_ots","Total_Readership_in_000s":"total_readership",
            "coScore_Index":"coscore_index","Total_Index":"total_index",
        },
        "ddl_extra":"edition TEXT, ad_rate REAL, article_length INTEGER, state TEXT, zone TEXT, total_ots REAL DEFAULT 0, total_readership REAL DEFAULT 0, coscore_index REAL DEFAULT 0, total_index REAL DEFAULT 0,",
        "genre_col":"genre","segment_col":"zone","segment_label":"Zone",
        "vol_col":"total_vol","read_col":"total_readership","idx_col":"coscore_index",
    },
    "online": {
        "label":"Online","icon":"🌐","table":"media_online",
        "unit_opts":["Index","Reach (Total OTS)","CCMs (Coverage)","Articles"],
        "unit_keys":["index","readership","coverage","articles"],
        "col_map":{
            "Company":"company","Article_Id":"article_id","Website":"source",
            "Publication_Date":"publication_date","Headline":"headline","Genre":"genre",
            "Program_Status":"program_status","Beneficial_Mentions":"beneficial_vol",
            "Neutral_Mentions":"neutral_vol","Adverse_Mentions":"adverse_vol",
            "Total_mentions":"total_vol","BN_mentions":"bn_vol",
            "B_Articles":"beneficial_art","N_Articles":"neutral_art",
            "A_Articles":"adverse_art","Total_Articles":"total_articles",
            "BN_Articles":"bn_articles","coScore":"coscore_index","coScore_Articles":"coscore_art",
        },
        "ddl_extra":"bn_vol INTEGER DEFAULT 0, bn_articles INTEGER DEFAULT 0, coscore_index REAL DEFAULT 0, coscore_art REAL DEFAULT 0, total_readership REAL DEFAULT 0, total_index REAL DEFAULT 0,",
        "genre_col":"genre","segment_col":"source","segment_label":"Website",
        "vol_col":"total_vol","read_col":"total_vol","idx_col":"coscore_index",
    },
    # Raw_Data3.xlsx columns → internal schema
    "tv": {
        "label":"TV","icon":"📺","table":"media_tv",
        "unit_opts":["Seconds","Clips","Index"],
        "unit_keys":["coverage","articles","index"],
        "col_map":{
            "Company":"company","Clip_Id":"article_id","Channel":"source",
            "Program_Date":"publication_date","Program":"genre",
            "Program_Telecast":"headline","Clip_Type":"program_status",
            "Time_Band":"zone","Article_Length":"article_length",
            "Total_seconds":"total_vol","Minutes":"total_readership",
            "Personality":"edition","Program_Month":"state","Client":"segment",
        },
        "ddl_extra":"zone TEXT, article_length INTEGER DEFAULT 0, total_readership REAL DEFAULT 0, coscore_index REAL DEFAULT 0, total_index REAL DEFAULT 0, edition TEXT, state TEXT, segment TEXT,",
        "genre_col":"genre","segment_col":"source","segment_label":"Channel",
        "vol_col":"total_vol","read_col":"total_readership","idx_col":"total_readership",
    },
    # Social_Media_ThirdParty_Data xlsx columns → internal schema
    "social": {
        "label":"Social Media","icon":"📱","table":"media_social",
        "unit_opts":["No. of Posts"],
        "unit_keys":["articles"],
        "col_map":{
            "Company":"company","Post ID":"article_id","Platforms":"source",
            "Post Date":"publication_date","Caption":"headline","Keyword":"genre",
            "Vertical":"program_status","Type Of Content":"zone",
            "Media Type":"edition","Profile Name":"state","Tonality":"segment",
            "Like Count":"total_readership","View Count":"total_vol",
            "Total Comment Count":"beneficial_art","Share Count":"neutral_art",
            "Positive Sentiment":"beneficial_vol","Neutral Sentiment":"neutral_vol",
            "Negative Sentiment":"adverse_vol","Hashtag Count":"article_length",
        },
        "ddl_extra":"zone TEXT, edition TEXT, state TEXT, segment TEXT, article_length INTEGER DEFAULT 0, total_readership REAL DEFAULT 0, coscore_index REAL DEFAULT 0, total_index REAL DEFAULT 0,",
        "genre_col":"genre","segment_col":"source","segment_label":"Platform",
        "vol_col":"total_vol","read_col":"total_readership","idx_col":"coscore_index",
    },
}


def _cfg(product_key: str, platform_key: str) -> dict:
    """Get platform config for given product. Falls back to HGEC PLATFORMS for backward compat."""
    if product_key and product_key in PRODUCTS:
        return PRODUCTS[product_key]["platforms"].get(platform_key, PLATFORMS.get(platform_key, {}))
    return PLATFORMS.get(platform_key, {})

BASE_DDL = """
CREATE TABLE IF NOT EXISTS {table} (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    company TEXT, article_id INTEGER, source TEXT,
    publication_date TEXT, headline TEXT, genre TEXT, program_status TEXT,
    beneficial_vol INTEGER DEFAULT 0, neutral_vol INTEGER DEFAULT 0,
    adverse_vol INTEGER DEFAULT 0, total_vol REAL DEFAULT 0,
    beneficial_art INTEGER DEFAULT 0, neutral_art INTEGER DEFAULT 0,
    adverse_art INTEGER DEFAULT 0, total_articles INTEGER DEFAULT 1,
    {extra}
    _platform TEXT
);
"""

def init_platform_db(product_key: str = None):
    """Create platform tables for the given product (or all products if None)."""
    products_to_init = [product_key] if product_key else list(PRODUCTS.keys())
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    required_cols = {"publication_date", "company", "total_articles",
                     "beneficial_art", "neutral_art", "adverse_art",
                     "beneficial_vol", "neutral_vol", "adverse_vol", "total_vol"}
    for prod_key in products_to_init:
        prod = PRODUCTS.get(prod_key, {})
        for plat_key, cfg in prod.get("platforms", {}).items():
            table = cfg["table"]
            existing = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table,)
            ).fetchone()
            if existing:
                present = {row[1] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()}
                if not required_cols.issubset(present):
                    conn.execute(f"DROP TABLE IF EXISTS {table}")
                    conn.commit()
            ddl = BASE_DDL.format(table=table, extra=cfg["ddl_extra"])
            conn.executescript(ddl)
    conn.commit()
    conn.close()

def load_platform_xlsx(xlsx_path, platform_key, product_key="hgec"):
    cfg = _cfg(product_key, platform_key)
    if not os.path.exists(xlsx_path): return 0
    df = pd.read_excel(xlsx_path)
    df.columns = [c.strip() for c in df.columns]
    df = df.rename(columns={k:v for k,v in cfg["col_map"].items() if k in df.columns})
    if "publication_date" in df.columns:
        df["publication_date"] = pd.to_datetime(df["publication_date"], errors="coerce").dt.strftime("%Y-%m-%d")
    defaults = {"beneficial_vol":0,"neutral_vol":0,"adverse_vol":0,"total_vol":0,
                "beneficial_art":0,"neutral_art":0,"adverse_art":0,"total_articles":1,
                "coscore_index":0,"total_readership":0,"total_index":0}
    for col, val in defaults.items():
        if col not in df.columns: df[col] = val
    df["total_articles"] = 1
    if platform_key == "social":
        likes    = pd.to_numeric(df.get("total_readership",0), errors="coerce").fillna(0)
        comments = pd.to_numeric(df.get("beneficial_art",0),   errors="coerce").fillna(0)
        shares   = pd.to_numeric(df.get("neutral_art",0),      errors="coerce").fillna(0)
        df["coscore_index"] = likes + comments + shares
    df["_platform"] = cfg["label"]
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.execute(f"PRAGMA table_info({cfg['table']})")
    db_cols = {row[1] for row in cursor.fetchall()}
    df = df[[c for c in df.columns if c in db_cols]]
    df.to_sql(cfg["table"], conn, if_exists="replace", index=False)
    conn.commit()
    n = conn.execute(f"SELECT COUNT(*) FROM {cfg['table']}").fetchone()[0]
    conn.close(); return n

def table_has_data(platform_key, product_key="hgec"):
    cfg = _cfg(product_key, platform_key)
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    try:
        exists = conn.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name=?",
                              (cfg["table"],)).fetchone()[0]
        if not exists: return False
        return conn.execute(f"SELECT COUNT(*) FROM {cfg['table']}").fetchone()[0] > 0
    finally: conn.close()

def get_platform_df(platform_key, product_key="hgec"):
    cfg = _cfg(product_key, platform_key)
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    try: return pd.read_sql_query(f"SELECT * FROM {cfg['table']}", conn)
    finally: conn.close()

def get_date_range(platform_key, product_key="hgec"):
    """Return (min_date, max_date) strings, or (None, None) if unavailable."""
    if not table_has_data(platform_key, product_key): return None, None
    cfg = _cfg(product_key, platform_key)
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    try:
        # Safety check: only query if column actually exists in this table
        cols = {row[1] for row in conn.execute(f"PRAGMA table_info({cfg['table']})").fetchall()}
        if "publication_date" not in cols:
            return None, None
        row = conn.execute(
            f"SELECT MIN(publication_date), MAX(publication_date) FROM {cfg['table']}"
        ).fetchone()
        return (row[0], row[1]) if row else (None, None)
    except Exception:
        return None, None
    finally:
        conn.close()

def _safe(df, col):
    if col in df.columns: return pd.to_numeric(df[col], errors="coerce").fillna(0)
    return pd.Series([0]*len(df), index=df.index)

def build_chart_data(df, platform_key, unit_key, product_key="hgec"):
    cfg = _cfg(product_key, platform_key)
    if df.empty: return _empty()
    df = df.copy()

    def get_metric(row):
        if unit_key=="articles":   return float(row.get("total_articles",1) or 1)
        if unit_key in ("readership","coverage"): return float(row.get("total_readership",0) or 0)
        if unit_key in ("index","coscore"):       return float(row.get("coscore_index",0) or 0)
        if unit_key=="mentions":   return float(row.get("total_vol",0) or 0)
        return float(row.get("total_articles",1) or 1)

    df["_metric"] = df.apply(get_metric, axis=1)
    companies = sorted(df["company"].dropna().unique().tolist()) if "company" in df.columns else []

    sov_total = max(_safe(df,"total_articles").sum(), 1)
    sov_data = []
    for co in companies:
        sub = df[df["company"]==co]
        arts = int(_safe(sub,"total_articles").sum())
        vols = float(_safe(sub,"total_vol").sum())
        sov_data.append({"company":co,"articles":arts,"volume":round(vols,1),"sov_pct":round((arts/sov_total)*100,1)})

    sentiment = []
    for co in companies:
        sub = df[df["company"]==co]
        tb  = float(_safe(sub,"beneficial_art").sum())
        tn  = float(_safe(sub,"neutral_art").sum())
        ta  = float(_safe(sub,"adverse_art").sum())
        tot = tb+tn+ta or 1
        sentiment.append({"company":co,"beneficial":round(tb/tot*100,1),"neutral":round(tn/tot*100,1),"adverse":round(ta/tot*100,1)})

    genre_data = []
    if cfg["genre_col"] in df.columns:
        grp = df.groupby(cfg["genre_col"])["_metric"].sum()
        tot = grp.sum() or 1
        genre_data = [{"genre":str(g),"value":float(v),"pct":round(v/tot*100,1)} for g,v in grp.nlargest(8).items()]

    seg_data = []
    if cfg["segment_col"] in df.columns:
        grp = df.groupby(cfg["segment_col"])["_metric"].sum().nlargest(8)
        tot = grp.sum() or 1
        seg_data = [{"name":str(k),"value":float(v),"pct":round(v/tot*100,1)} for k,v in grp.items()]

    trend_dates, trend_series = [], {}
    if "publication_date" in df.columns:
        trend_dates = sorted(df["publication_date"].dropna().unique().tolist())
        for co in companies:
            sub = df[df["company"]==co].groupby("publication_date")["_metric"].sum()
            trend_series[co] = [float(sub.get(d,0)) for d in trend_dates]

    all_dates = trend_dates
    tv_extra = {}
    if platform_key=="tv":
        if "zone" in df.columns:
            tv_extra["timebands"] = {str(k):float(v) for k,v in df.groupby("zone")["_metric"].sum().items()}
        if "program_status" in df.columns:
            tv_extra["clip_types"] = {str(k):float(v) for k,v in df.groupby("program_status")["_metric"].sum().items()}

    return {
        "platform":cfg["label"],"unit_key":unit_key,
        "unit_label":dict(zip(cfg["unit_keys"],cfg["unit_opts"])).get(unit_key,"Count"),
        "kpi":{"index":round(float(_safe(df,"coscore_index").sum()),1),
               "articles":int(_safe(df,"total_articles").sum()),
               "reach":round(float(_safe(df,"total_readership").sum()),1),
               "vol":round(float(_safe(df,"total_vol").sum()),1)},
        "companies":companies,"sov":sov_data,"sentiment":sentiment,
        "genre":genre_data,
        "segment":{"label":cfg["segment_label"],"data":seg_data},
        "trend":{"dates":trend_dates,"series":trend_series},
        "date_range":{"min":all_dates[0] if all_dates else "","max":all_dates[-1] if all_dates else "","all":all_dates},
        "unit_opts":list(zip(cfg["unit_keys"],cfg["unit_opts"])),
        "tv_extra":tv_extra,"platform_key":platform_key,
    }

def _empty():
    return {"platform":"","unit_key":"articles","unit_label":"Count",
            "kpi":{"index":0,"articles":0,"reach":0,"vol":0},
            "companies":[],"sov":[],"sentiment":[],"genre":[],
            "segment":{"label":"Segment","data":[]},"trend":{"dates":[],"series":{}},
            "date_range":{"min":"","max":"","all":[]},"unit_opts":[("articles","Count")],
            "tv_extra":{},"platform_key":""}

def _build_html(chart_data, platform_key, product_key="hgec"):
    cfg = _cfg(product_key, platform_key)
    data_json = json.dumps(chart_data)
    plat_color = {"print":"#2563EB","online":"#059669","tv":"#7C3AED","social":"#DB2777"}.get(platform_key,"#4F46E5")
    plat_light = {"print":"#DBEAFE","online":"#D1FAE5","tv":"#EDE9FE","social":"#FCE7F3"}.get(platform_key,"#EEF2FF")
    is_tv     = "true" if platform_key=="tv"     else "false"
    is_social = "true" if platform_key=="social" else "false"
    kpi1_lbl  = "Total Air Time (min)" if platform_key=="tv" else ("Total Views" if platform_key=="social" else "Growth in Index Score")
    kpi1_sub  = "minutes across all clips" if platform_key=="tv" else ("total view count" if platform_key=="social" else "coScore Index")
    kpi2_lbl  = "Total Clips" if platform_key=="tv" else ("Total Posts" if platform_key=="social" else "Articles")
    kpi3_lbl  = "Total Likes" if platform_key=="social" else "Reach & Visibility"

    return f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#F8FAFF;color:#1E293B;font-size:13px}}
.filters{{background:#fff;border-bottom:1px solid #E2E8F0;padding:10px 16px;display:flex;flex-wrap:wrap;gap:10px;align-items:flex-end;box-shadow:0 1px 3px rgba(0,0,0,.06)}}
.fg{{display:flex;flex-direction:column;gap:3px}}
.fg label{{font-size:10px;color:#64748B;font-weight:600;text-transform:uppercase;letter-spacing:.05em}}
.fg select,.fg input{{padding:5px 8px;font-size:12px;border:1px solid #CBD5E1;border-radius:6px;background:#F8FAFF;color:#1E293B;height:30px;min-width:110px;outline:none}}
.fg input[type=date]{{min-width:128px}}
.export-btn{{padding:0 14px;height:30px;background:{plat_color};color:#fff;border:none;border-radius:6px;font-size:12px;font-weight:500;cursor:pointer;margin-top:17px}}
.export-btn:hover{{opacity:.85}}
.kpi-row{{display:grid;grid-template-columns:repeat(3,1fr);gap:10px;padding:12px 16px}}
.kpi{{background:#fff;border:1px solid #E2E8F0;border-radius:10px;padding:12px 14px;box-shadow:0 1px 3px rgba(0,0,0,.04)}}
.kpi.primary{{background:{plat_light};border-color:{plat_color}55}}
.kpi-lbl{{font-size:10px;color:#64748B;margin-bottom:4px;font-weight:600;text-transform:uppercase;letter-spacing:.05em}}
.kpi.primary .kpi-lbl{{color:{plat_color}}}
.kpi-val{{font-size:22px;font-weight:700;line-height:1.1;color:#1E293B}}
.kpi.primary .kpi-val{{color:{plat_color}}}
.kpi-badge{{display:inline-flex;align-items:center;gap:3px;font-size:10px;margin-top:4px;padding:2px 7px;border-radius:20px;background:#DCFCE7;color:#166534}}
.kpi-sub{{font-size:10px;color:#94A3B8;margin-top:3px}}
.grid-2{{display:grid;grid-template-columns:1fr 1fr;gap:10px;padding:0 16px 10px}}
.panel{{background:#fff;border:1px solid #E2E8F0;border-radius:10px;padding:14px 15px;box-shadow:0 1px 3px rgba(0,0,0,.04)}}
.pt{{font-size:12px;font-weight:600;color:#1E293B;margin-bottom:10px;display:flex;justify-content:space-between;align-items:center}}
.legend{{display:flex;flex-wrap:wrap;gap:8px;margin-bottom:8px}}
.li{{display:flex;align-items:center;gap:4px;font-size:10px;color:#64748B}}
.ld{{width:9px;height:9px;border-radius:2px;flex-shrink:0}}
.split{{display:grid;grid-template-columns:1fr 150px;gap:10px;align-items:center}}
canvas{{display:block}}
.ttog{{display:flex;border:1px solid #CBD5E1;border-radius:5px;overflow:hidden}}
.tt{{padding:3px 8px;font-size:10px;background:none;border:none;cursor:pointer;color:#64748B;border-right:1px solid #CBD5E1}}
.tt:last-child{{border-right:none}}
.tt.active{{background:{plat_color};color:#fff}}
.dt{{width:100%;border-collapse:collapse;font-size:11px}}
.dt th{{text-align:left;padding:4px 5px;color:#94A3B8;font-weight:600;border-bottom:1px solid #E2E8F0;font-size:10px}}
.dt td{{padding:4px 5px;border-bottom:1px solid #F1F5F9;color:#475569}}
.dt tr:last-child td{{border-bottom:none}}
.sov-bar{{height:6px;border-radius:3px;background:{plat_color};margin-top:3px;transition:width .4s}}
</style></head><body>

<div class="filters">
  <div class="fg"><label>Start Date</label><input type="date" id="f-start" onchange="applyFilters()"></div>
  <div class="fg"><label>End Date</label><input type="date" id="f-end" onchange="applyFilters()"></div>
  <div class="fg"><label>Company</label><select id="f-company" onchange="applyFilters()"><option value="all">All Companies</option></select></div>
  <div class="fg"><label>Unit</label><select id="f-unit" onchange="applyFilters()"></select></div>
  <button class="export-btn" onclick="doExport()">⬇ Export CSV</button>
</div>

<div class="kpi-row">
  <div class="kpi primary">
    <div class="kpi-lbl">{kpi1_lbl}</div>
    <div class="kpi-val" id="kpi-index">—</div>
    <div class="kpi-badge" id="kpi-trend">▲ live</div>
    <div class="kpi-sub">{kpi1_sub}</div>
  </div>
  <div class="kpi">
    <div class="kpi-lbl">{kpi2_lbl}</div>
    <div class="kpi-val" id="kpi-articles">—</div>
    <div class="kpi-sub">total in period</div>
  </div>
  <div class="kpi">
    <div class="kpi-lbl">{kpi3_lbl}</div>
    <div class="kpi-val" id="kpi-reach">—</div>
    <div class="kpi-sub" id="kpi-reach-sub">aggregate</div>
  </div>
</div>

<div class="grid-2">
  <div class="panel">
    <div class="pt">Share of Voice (SOV)</div>
    <div style="position:relative;width:100%;height:170px"><canvas id="sovChart"></canvas></div>
  </div>
  <div class="panel">
    <div class="pt">Sentiment / Clip Type</div>
    <div class="legend">
      <div class="li"><div class="ld" style="background:#16A34A"></div>Beneficial / Special Feature</div>
      <div class="li"><div class="ld" style="background:#6B7280"></div>Neutral / News</div>
      <div class="li"><div class="ld" style="background:#DC2626"></div>Adverse</div>
    </div>
    <div style="position:relative;width:100%;height:145px"><canvas id="sentChart"></canvas></div>
  </div>
</div>

<div class="grid-2">
  <div class="panel">
    <div class="pt">Spread by Genre / Keyword / Programme</div>
    <div class="split">
      <div><div style="position:relative;width:100%;height:165px"><canvas id="genreChart"></canvas></div></div>
      <table class="dt">
        <thead><tr><th>Item</th><th>%</th></tr></thead>
        <tbody id="genreTbl"></tbody>
      </table>
    </div>
  </div>
  <div class="panel">
    <div class="pt">
      Trend Analysis
      <div class="ttog">
        <button class="tt active" onclick="setGroup('day',this)">Day</button>
        <button class="tt" onclick="setGroup('15d',this)">15-Days</button>
        <button class="tt" onclick="setGroup('monthly',this)">Monthly</button>
        <button class="tt" onclick="setGroup('quarterly',this)">Quarterly</button>
      </div>
    </div>
    <div class="legend" id="trendLeg"></div>
    <div style="position:relative;width:100%;height:155px"><canvas id="trendChart"></canvas></div>
  </div>
</div>

<div style="padding:0 16px 16px">
  <div class="panel">
    <div class="pt">Company Detail</div>
    <table class="dt" style="width:100%">
      <thead><tr><th>Company</th><th>Count</th><th>SOV %</th><th style="width:35%">Share</th><th>Volume</th></tr></thead>
      <tbody id="sovTbl"></tbody>
    </table>
  </div>
</div>

<script src="https://cdnjs.cloudflare.com/ajax/libs/Chart.js/4.4.1/chart.umd.js"></script>
<script>
const SD={data_json};
const PC="{plat_color}";
const IS_TV={is_tv};
const IS_SOCIAL={is_social};
const CO_COLORS=["#2563EB","#059669","#D97706","#7C3AED","#DB2777","#0891B2","#65A30D","#9333EA"];
const GENRE_COLORS=["#4F46E5","#059669","#D97706","#DB2777","#0891B2","#65A30D","#9333EA","#EF4444"];
const GRID_C="rgba(203,213,225,0.5)";const TEXT_C="#64748B";
let charts={{}};let activeGroup="day";

function fmt(n){{n=Number(n)||0;if(Math.abs(n)>=1e6)return(n/1e6).toFixed(1)+"M";if(Math.abs(n)>=1e3)return(n/1e3).toFixed(1)+"K";return Math.round(n).toString();}}

function initUI(){{
  const sel=document.getElementById("f-company");
  SD.companies.forEach(c=>{{const o=document.createElement("option");o.value=c;o.textContent=c;sel.appendChild(o);}});
  const uSel=document.getElementById("f-unit");
  SD.unit_opts.forEach(([k,l])=>{{const o=document.createElement("option");o.value=k;o.textContent=l;uSel.appendChild(o);}});
  const allDates=SD.date_range.all;
  const startEl=document.getElementById("f-start");const endEl=document.getElementById("f-end");
  if(allDates.length){{
    startEl.value=allDates[0];startEl.min=allDates[0];startEl.max=allDates[allDates.length-1];
    endEl.value=allDates[allDates.length-1];endEl.min=allDates[0];endEl.max=allDates[allDates.length-1];
  }}
}}

function filterData(){{
  const co=document.getElementById("f-company").value;
  const start=document.getElementById("f-start").value;
  const end=document.getElementById("f-end").value;
  let sov=SD.sov,sent=SD.sentiment,genre=SD.genre;
  let trendSeries=SD.trend.series,trendDates=SD.trend.dates;
  const dateIdxs=trendDates.map((d,i)=>d>=start&&d<=end?i:-1).filter(i=>i>=0);
  const filtDates=dateIdxs.map(i=>trendDates[i]);
  let filtSeries={{}};
  if(co==="all"){{Object.keys(trendSeries).forEach(c=>{{filtSeries[c]=dateIdxs.map(i=>trendSeries[c][i]||0);}});}}
  else{{if(trendSeries[co])filtSeries[co]=dateIdxs.map(i=>trendSeries[co][i]||0);sov=sov.filter(r=>r.company===co);sent=sent.filter(r=>r.company===co);}}
  return{{sov,sent,genre,trend:{{dates:filtDates,series:filtSeries}},kpi:SD.kpi}};
}}

function updateKPIs(d){{
  if(IS_TV){{document.getElementById("kpi-index").textContent=fmt(SD.kpi.reach);}}
  else if(IS_SOCIAL){{document.getElementById("kpi-index").textContent=fmt(SD.kpi.vol);}}
  else{{document.getElementById("kpi-index").textContent=fmt(SD.kpi.index);}}
  document.getElementById("kpi-articles").textContent=fmt(SD.kpi.articles);
  document.getElementById("kpi-reach").textContent=fmt(SD.kpi.reach);
  document.getElementById("kpi-trend").textContent="▲ live data";
}}

function destroyChart(id){{if(charts[id]){{charts[id].destroy();delete charts[id];}}}}

function updateSOV(d){{
  if(!d.sov.length)return;
  const labels=d.sov.map(r=>r.company);
  destroyChart("sovChart");
  charts.sovChart=new Chart(document.getElementById("sovChart"),{{
    type:"bar",
    data:{{labels,datasets:[
      {{label:"Count",data:d.sov.map(r=>r.articles),backgroundColor:PC+"CC",borderRadius:4,borderWidth:0}},
      {{label:"Volume",data:d.sov.map(r=>r.volume),backgroundColor:PC+"44",borderRadius:4,borderWidth:0}},
    ]}},
    options:{{responsive:true,maintainAspectRatio:false,
      plugins:{{legend:{{labels:{{color:TEXT_C,font:{{size:10}},boxWidth:10,padding:10}}}},tooltip:{{callbacks:{{label:c=>c.dataset.label+": "+fmt(c.raw)}}}}}},
      scales:{{x:{{grid:{{display:false}},ticks:{{color:TEXT_C,font:{{size:10}}}}}},y:{{grid:{{color:GRID_C}},ticks:{{color:TEXT_C,font:{{size:10}},callback:v=>fmt(v)}}}}}}
    }}
  }});
  const totArt=d.sov.reduce((s,r)=>s+r.articles,0)||1;
  document.getElementById("sovTbl").innerHTML=d.sov.map(r=>{{
    const pct=((r.articles/totArt)*100).toFixed(1);const barW=Math.round(r.articles/totArt*100);
    return`<tr><td style="font-weight:600;color:#1E293B">${{r.company}}</td><td>${{fmt(r.articles)}}</td><td style="color:${{PC}};font-weight:600">${{pct}}%</td><td><div style="background:#F1F5F9;border-radius:3px;height:6px;width:100%"><div class="sov-bar" style="width:${{barW}}%"></div></div></td><td>${{fmt(r.volume)}}</td></tr>`;
  }}).join("");
}}

function updateSentiment(d){{
  destroyChart("sentChart");
  if(IS_TV){{
    const ctypes=SD.tv_extra.clip_types||{{}};
    const labels=Object.keys(ctypes);const vals=Object.values(ctypes);
    if(!labels.length)return;
    charts.sentChart=new Chart(document.getElementById("sentChart"),{{
      type:"doughnut",
      data:{{labels,datasets:[{{data:vals,backgroundColor:["#16A34A","#6B7280","#2563EB","#D97706"],borderWidth:2,borderColor:"#F8FAFF"}}]}},
      options:{{responsive:true,maintainAspectRatio:false,cutout:"55%",plugins:{{legend:{{labels:{{color:TEXT_C,font:{{size:10}},boxWidth:8,padding:6}}}}}}}}
    }});
  }}else{{
    if(!d.sent.length)return;
    charts.sentChart=new Chart(document.getElementById("sentChart"),{{
      type:"bar",
      data:{{labels:d.sent.map(r=>r.company),datasets:[
        {{label:"Beneficial",data:d.sent.map(r=>r.beneficial),backgroundColor:"#16A34A",borderWidth:0,stack:"s"}},
        {{label:"Neutral",data:d.sent.map(r=>r.neutral),backgroundColor:"#6B7280",borderWidth:0,stack:"s"}},
        {{label:"Adverse",data:d.sent.map(r=>r.adverse),backgroundColor:"#DC2626",borderWidth:0,stack:"s"}},
      ]}},
      options:{{responsive:true,maintainAspectRatio:false,
        plugins:{{legend:{{display:false}},tooltip:{{callbacks:{{label:c=>c.dataset.label+": "+c.raw+"%"}}}}}},
        scales:{{x:{{stacked:true,grid:{{display:false}},ticks:{{color:TEXT_C,font:{{size:10}}}}}},y:{{stacked:true,max:100,ticks:{{callback:v=>v+"%",color:TEXT_C,font:{{size:10}}}},grid:{{color:GRID_C}}}}}}
      }}
    }});
  }}
}}

function updateGenre(d){{
  if(!d.genre.length){{destroyChart("genreChart");return;}}
  const labels=d.genre.map(g=>g.genre);const vals=d.genre.map(g=>g.value);const total=vals.reduce((a,b)=>a+b,0)||1;
  document.getElementById("genreTbl").innerHTML=d.genre.map((g,i)=>
    `<tr><td style="display:flex;align-items:center;gap:4px"><div style="width:7px;height:7px;border-radius:2px;flex-shrink:0;background:${{GENRE_COLORS[i%GENRE_COLORS.length]}}"></div>${{g.genre.length>16?g.genre.slice(0,14)+"…":g.genre}}</td><td>${{g.pct}}%</td></tr>`
  ).join("");
  destroyChart("genreChart");
  charts.genreChart=new Chart(document.getElementById("genreChart"),{{
    type:"doughnut",
    data:{{labels,datasets:[{{data:vals,backgroundColor:GENRE_COLORS.slice(0,labels.length),borderWidth:2,borderColor:"#F8FAFF"}}]}},
    options:{{responsive:true,maintainAspectRatio:false,cutout:"58%",plugins:{{legend:{{display:false}},tooltip:{{callbacks:{{label:c=>c.label+": "+((c.raw/total)*100).toFixed(1)+"%"}}}}}}}}
  }});
}}

function groupTrend(dates,series,grouping){{
  const grouped={{}};
  dates.forEach((d,i)=>{{
    let key;const dt=new Date(d);
    if(grouping==="day")key=d;
    else if(grouping==="15d")key=d.slice(0,8)+(dt.getDate()<=15?"01":"16");
    else if(grouping==="monthly")key=d.slice(0,7);
    else key=d.slice(0,4)+"-Q"+Math.ceil((dt.getMonth()+1)/3);
    if(!grouped[key])grouped[key]={{}};
    Object.keys(series).forEach(co=>{{grouped[key][co]=(grouped[key][co]||0)+(series[co][i]||0);}});
  }});
  return grouped;
}}

function updateTrend(d){{
  const grouped=groupTrend(d.trend.dates,d.trend.series,activeGroup);
  const keys=Object.keys(grouped).sort();const cos=Object.keys(d.trend.series);
  document.getElementById("trendLeg").innerHTML=cos.map((c,i)=>
    `<div class="li"><div class="ld" style="background:${{CO_COLORS[i%CO_COLORS.length]}}"></div>${{c}}</div>`
  ).join("");
  const datasets=cos.map((c,i)=>({{label:c,data:keys.map(k=>grouped[k][c]||0),
    borderColor:CO_COLORS[i%CO_COLORS.length],backgroundColor:CO_COLORS[i%CO_COLORS.length]+"22",
    borderWidth:2,pointRadius:3,tension:.3,fill:false}}));
  destroyChart("trendChart");
  charts.trendChart=new Chart(document.getElementById("trendChart"),{{
    type:"line",data:{{labels:keys.map(k=>k.length>7?k.slice(5):k),datasets}},
    options:{{responsive:true,maintainAspectRatio:false,plugins:{{legend:{{display:false}}}},
      scales:{{x:{{grid:{{display:false}},ticks:{{color:TEXT_C,font:{{size:10}},maxTicksLimit:12}}}},y:{{grid:{{color:GRID_C}},ticks:{{color:TEXT_C,font:{{size:10}},callback:v=>fmt(v)}}}}}}
    }}
  }});
}}

function applyFilters(){{
  const filtered=filterData();
  updateKPIs(filtered);updateSOV(filtered);updateSentiment(filtered);updateGenre(filtered);updateTrend(filtered);
}}

function setGroup(g,btn){{
  activeGroup=g;document.querySelectorAll(".tt").forEach(b=>b.classList.remove("active"));btn.classList.add("active");updateTrend(filterData());
}}

function doExport(){{window.parent.postMessage({{type:"cbc_export",platform:"{platform_key}"}}, "*");}}

initUI();applyFilters();
</script></body></html>"""

def render_platform_dashboard(username, role, platform_key="print", product_key="hgec"):
    cfg = _cfg(product_key, platform_key)
    is_admin = (role=="admin")
    _icon = cfg.get('icon', '') if cfg.get('icon', '') not in ('print','online','tv','social','') else {'print':'🗞️','online':'🌐','tv':'📺','social':'📱'}.get(cfg.get('icon',''), '📊')
    st.markdown(f"## {_icon} {cfg['label']} — Media Intelligence")

    if is_admin:
        with st.expander(f"⚙️ Data Management — {cfg['label']} (Admin)", expanded=not table_has_data(platform_key)):
            st.caption(f"Upload your {cfg['label']} Raw Data (.xlsx)")
            c1,c2 = st.columns([3,1])
            with c1:
                uploaded = st.file_uploader("Upload xlsx",type=["xlsx"],key=f"upload_{platform_key}_{product_key}",label_visibility="collapsed")
            with c2:
                st.markdown("<div style='height:28px'></div>",unsafe_allow_html=True)
                if uploaded and st.button("Load Data",key=f"load_{platform_key}_{product_key}"):
                    with tempfile.NamedTemporaryFile(delete=False,suffix=".xlsx") as tmp:
                        tmp.write(uploaded.read()); tmp_path=tmp.name
                    with st.spinner("Loading…"):
                        n=load_platform_xlsx(tmp_path,platform_key,product_key)
                    os.unlink(tmp_path); st.success(f"✅ {n} rows loaded."); st.rerun()
            if table_has_data(platform_key, product_key):
                conn=sqlite3.connect(DB_PATH,check_same_thread=False)
                n=conn.execute(f"SELECT COUNT(*) FROM {cfg['table']}").fetchone()[0]
                cos=conn.execute(f"SELECT DISTINCT company FROM {cfg['table']} ORDER BY company").fetchall()
                mn,mx=get_date_range(platform_key, product_key); conn.close()
                st.info(f"**{n} rows** · Companies: {', '.join(r[0] for r in cos)} · Dates: {mn} → {mx}")

    if not table_has_data(platform_key, product_key):
        st.warning(f"No {cfg['label']} data loaded. " + ("Upload via Data Management above." if is_admin else "Ask an admin to upload the data."))
        return

    df = get_platform_df(platform_key, product_key)
    chart_data = build_chart_data(df, platform_key, "articles", product_key)
    html = _build_html(chart_data, platform_key, product_key)
    st.components.v1.html(html, height=980, scrolling=False)

    st.markdown("---")
    _render_export_section(df, chart_data, platform_key, cfg, role)

# ── Chart Export (PDF & Word — charts only, no raw data) ─────────────────────

def _build_chart_images(chart_data: dict, platform_key: str) -> list:
    """Render chart_data into a list of (title, PNG-bytes) using matplotlib."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np

    COLORS = ["#4F46E5","#059669","#D97706","#DC2626","#0891B2","#7C3AED"]
    SENT_COLORS = {"Beneficial":"#16A34A","Neutral":"#6B7280","Adverse":"#DC2626"}
    images = []

    def save_fig(fig, title):
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=140, bbox_inches="tight",
                    facecolor="#FFFFFF", edgecolor="none")
        plt.close(fig)
        buf.seek(0)
        images.append((title, buf.read()))

    def fmt_num(v):
        if v >= 1_000_000: return f"{v/1_000_000:.1f}M"
        if v >= 1_000: return f"{v/1_000:.1f}K"
        return str(int(v))

    # ── 1. SOV Bar Chart ────────────────────────────────────────
    sov = chart_data.get("sov", [])
    if sov:
        companies = [r["company"] for r in sov]
        articles  = [r["articles"] for r in sov]
        fig, ax = plt.subplots(figsize=(9, 4))
        bars = ax.bar(companies, articles,
                      color=[COLORS[i % len(COLORS)] for i in range(len(companies))],
                      width=0.55, zorder=3)
        for bar, val in zip(bars, articles):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(articles)*0.01,
                    fmt_num(val), ha="center", va="bottom", fontsize=9, fontweight="600")
        ax.set_title("Share of Voice — Article Count", fontsize=12, fontweight="700", pad=12)
        ax.set_ylabel("Articles", fontsize=9)
        ax.set_facecolor("#F8FAFF"); fig.patch.set_facecolor("#FFFFFF")
        ax.grid(axis="y", color="#E2E8F0", zorder=0)
        ax.spines[["top","right","left"]].set_visible(False)
        ax.tick_params(axis="x", labelsize=9); ax.tick_params(axis="y", labelsize=8)
        fig.tight_layout()
        save_fig(fig, "Share of Voice")

    # ── 2. Sentiment Stacked Bar ────────────────────────────────
    sent = chart_data.get("sent", [])
    if sent:
        companies = [r["company"] for r in sent]
        ben = [r["beneficial"] for r in sent]
        neu = [r["neutral"]    for r in sent]
        adv = [r["adverse"]    for r in sent]
        x   = np.arange(len(companies))
        fig, ax = plt.subplots(figsize=(9, 4))
        b1 = ax.bar(x, ben, label="Beneficial", color="#16A34A", width=0.55, zorder=3)
        b2 = ax.bar(x, neu, label="Neutral",    color="#6B7280", bottom=ben, width=0.55, zorder=3)
        b3 = ax.bar(x, adv, label="Adverse",    color="#DC2626",
                    bottom=[b+n for b,n in zip(ben,neu)], width=0.55, zorder=3)
        ax.set_xticks(x); ax.set_xticklabels(companies, fontsize=9)
        ax.set_title("Sentiment Distribution (%)", fontsize=12, fontweight="700", pad=12)
        ax.set_ylabel("Share (%)", fontsize=9)
        ax.set_facecolor("#F8FAFF"); fig.patch.set_facecolor("#FFFFFF")
        ax.grid(axis="y", color="#E2E8F0", zorder=0)
        ax.spines[["top","right","left"]].set_visible(False)
        ax.legend(fontsize=8, frameon=False)
        ax.tick_params(axis="y", labelsize=8)
        fig.tight_layout()
        save_fig(fig, "Sentiment Distribution")

    # ── 3. Genre / Segment Donut ────────────────────────────────
    genre = chart_data.get("genre", [])
    if genre:
        labels = [g["genre"] for g in genre]
        vals   = [g["value"] for g in genre]
        fig, ax = plt.subplots(figsize=(6, 5))
        wedge_colors = [COLORS[i % len(COLORS)] for i in range(len(labels))]
        wedges, texts, autotexts = ax.pie(
            vals, labels=None, autopct="%1.1f%%",
            colors=wedge_colors, startangle=90,
            pctdistance=0.75, wedgeprops=dict(width=0.55, edgecolor="white", linewidth=2)
        )
        for at in autotexts:
            at.set_fontsize(8); at.set_color("white"); at.set_fontweight("600")
        ax.legend(wedges, [l[:20] for l in labels], loc="lower center",
                  bbox_to_anchor=(0.5, -0.12), ncol=3, fontsize=7, frameon=False)
        ax.set_title("Genre / Segment Breakdown", fontsize=12, fontweight="700", pad=12)
        fig.tight_layout()
        save_fig(fig, "Genre / Segment Breakdown")

    # ── 4. Trend Line ───────────────────────────────────────────
    trend = chart_data.get("trend", {})
    dates  = trend.get("dates", [])
    series = trend.get("series", {})
    if dates and series:
        fig, ax = plt.subplots(figsize=(9, 4))
        for i, (company, vals) in enumerate(series.items()):
            ax.plot(range(len(dates)), vals,
                    label=company, color=COLORS[i % len(COLORS)],
                    linewidth=2, marker="o", markersize=3)
        step = max(1, len(dates)//10)
        ax.set_xticks(range(0, len(dates), step))
        ax.set_xticklabels([d[5:] for d in dates[::step]], rotation=30, fontsize=7, ha="right")
        ax.set_title("Coverage Trend Over Time", fontsize=12, fontweight="700", pad=12)
        ax.set_ylabel("Volume", fontsize=9)
        ax.set_facecolor("#F8FAFF"); fig.patch.set_facecolor("#FFFFFF")
        ax.grid(color="#E2E8F0", zorder=0)
        ax.spines[["top","right"]].set_visible(False)
        ax.legend(fontsize=8, frameon=False)
        ax.tick_params(axis="y", labelsize=8)
        fig.tight_layout()
        save_fig(fig, "Coverage Trend")

    return images


def _generate_pdf(images: list, platform_label: str) -> bytes:
    """Build a PDF with one chart per page using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Image as RLImage, Spacer, Paragraph, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=1.8*cm, bottomMargin=1.8*cm,
                            leftMargin=1.5*cm, rightMargin=1.5*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                 fontSize=16, textColor=colors.HexColor("#4F46E5"),
                                 spaceAfter=4, alignment=TA_CENTER)
    sub_style = ParagraphStyle("Sub", parent=styles["Normal"],
                               fontSize=9, textColor=colors.HexColor("#64748B"),
                               spaceAfter=16, alignment=TA_CENTER)
    chart_title_style = ParagraphStyle("ChartTitle", parent=styles["Heading2"],
                                       fontSize=11, textColor=colors.HexColor("#1E293B"),
                                       spaceBefore=12, spaceAfter=6)

    story = []
    story.append(Paragraph(f"TAM — EIKONA Dashboard", title_style))
    story.append(Paragraph(f"{platform_label} Media Intelligence Report · {datetime.now().strftime('%d %B %Y')}", sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E2E8F0"), spaceAfter=18))

    page_w = A4[0] - 3*cm   # usable width
    for chart_title, png_bytes in images:
        story.append(Paragraph(chart_title, chart_title_style))
        img_buf = io.BytesIO(png_bytes)
        img = RLImage(img_buf)
        # Scale to fit width, preserve aspect ratio
        ratio = img.drawHeight / img.drawWidth
        img.drawWidth  = page_w
        img.drawHeight = page_w * ratio
        story.append(img)
        story.append(Spacer(1, 0.4*cm))

    doc.build(story)
    buf.seek(0)
    return buf.read()


def _generate_word(images: list, platform_label: str) -> bytes:
    """Build a Word doc with one chart per page using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = Document()

    # Page margins
    for section in document.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Title
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("TAM — EIKONA Dashboard")
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    sub_p = document.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        f"{platform_label} Media Intelligence Report  ·  {datetime.now().strftime('%d %B %Y')}")
    sub_run.font.size = Pt(9); sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    document.add_paragraph()  # spacer

    for chart_title, png_bytes in images:
        # Chart heading
        h = document.add_heading(chart_title, level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Inline image
        img_buf = io.BytesIO(png_bytes)
        document.add_picture(img_buf, width=Inches(6.0))
        last_para = document.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()  # spacer between charts

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()


def _render_export_section(df, chart_data, platform_key, cfg, role):
    """Render export buttons: PDF & Word (charts only) for all; CSV only for admin/partner."""
    label = cfg["label"]
    row_count = len(df)
    company_count = df["company"].nunique() if "company" in df.columns else 0

    st.markdown(
        f"<p style='font-size:.78rem;color:#64748B;margin-bottom:8px'>"
        f"📊 {row_count:,} records · {company_count} companies</p>",
        unsafe_allow_html=True
    )

    # Session keys for generated report bytes
    pdf_key  = f"_report_pdf_{platform_key}"
    word_key = f"_report_word_{platform_key}"

    col1, col2, col3 = st.columns([1, 1, 1])

    with col1:
        if st.button("📄 Generate PDF Report", key=f"gen_pdf_{platform_key}",
                     help="Charts only — no raw data"):
            with st.spinner("Building PDF…"):
                imgs = _build_chart_images(chart_data, platform_key)
                st.session_state[pdf_key] = _generate_pdf(imgs, label)

    with col2:
        if st.button("📝 Generate Word Report", key=f"gen_word_{platform_key}",
                     help="Charts only — no raw data"):
            with st.spinner("Building Word doc…"):
                imgs = _build_chart_images(chart_data, platform_key)
                st.session_state[word_key] = _generate_word(imgs, label)

    with col3:
        # CSV download only for admin and partner — NOT for client
        if role in ("admin", "partner"):
            st.download_button(
                label="⬇ Download Raw CSV",
                data=df.to_csv(index=False).encode(),
                file_name=f"{platform_key}_export.csv",
                mime="text/csv",
                key=f"dl_csv_{platform_key}",
                help="Raw data export (Admin & Partner only)"
            )

    # ── Download zone — appears below, auto-scrolls into view when ready ──
    pdf_ready  = pdf_key  in st.session_state
    word_ready = word_key in st.session_state

    if pdf_ready or word_ready:
        anchor_id = f"dl_zone_{platform_key}"
        # Inject named anchor + JS scroll
        st.markdown(
            f'''<div id="{anchor_id}" style="margin-top:12px"></div>
<script>
(function(){{
    var el = window.parent.document.getElementById("{anchor_id}");
    if(!el) el = document.getElementById("{anchor_id}");
    if(el) {{ setTimeout(function(){{ el.scrollIntoView({{behavior:"smooth",block:"start"}}); }}, 120); }}
    // fallback: scroll parent window to bottom of main area
    setTimeout(function(){{
        var frames = window.parent.document.querySelectorAll("iframe");
        frames.forEach(function(f){{ try{{ f.contentWindow.scrollTo({{top:99999,behavior:"smooth"}}); }}catch(e){{}} }});
        window.parent.scrollTo({{top: window.parent.document.body.scrollHeight, behavior:"smooth"}});
    }}, 150);
}})();
</script>''',
            unsafe_allow_html=True
        )

        st.markdown(
            "<div style='background:#F0FDF4;border:1px solid #BBF7D0;border-radius:10px;"
            "padding:14px 20px;margin-top:4px'>",
            unsafe_allow_html=True
        )
        st.markdown("**✅ Your report is ready — click below to download:**")
        dl_c1, dl_c2, dl_c3 = st.columns([1, 1, 1])

        with dl_c1:
            if pdf_ready:
                st.download_button(
                    label="⬇ Download PDF",
                    data=st.session_state[pdf_key],
                    file_name=f"{platform_key}_charts_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf",
                    key=f"dl_pdf_{platform_key}"
                )
        with dl_c2:
            if word_ready:
                st.download_button(
                    label="⬇ Download Word",
                    data=st.session_state[word_key],
                    file_name=f"{platform_key}_charts_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_word_{platform_key}"
                )
        with dl_c3:
            if st.button("🗑 Clear", key=f"clear_reports_{platform_key}", help="Clear generated reports"):
                st.session_state.pop(pdf_key, None)
                st.session_state.pop(word_key, None)
                st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)